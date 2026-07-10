import os
import time
import logging
from typing import Iterable

from docx import Document as DocxDocument
from fastapi import UploadFile
from pypdf import PdfReader

from ..database import models
from ..database.database import SessionLocal
from ..llm.openai_service import OpenAIService
from ..rag.chunking import split_text

logger = logging.getLogger(__name__)


class UploadService:
    def __init__(self, db: SessionLocal):
        self.openai_service = OpenAIService()
        self.db = db

    def save_file(self, user_id: int, upload: UploadFile) -> models.Document:
        upload_dir = os.path.join(os.path.dirname(__file__), "..", "..", "uploads")
        os.makedirs(upload_dir, exist_ok=True)
        path = os.path.join(upload_dir, upload.filename)

        # Seek to beginning of file before reading
        upload.file.seek(0)
        with open(path, "wb") as file_obj:
            file_obj.write(upload.file.read())

        content = self._extract_content(path, upload.filename)
        if not content or not content.strip():
            raise ValueError("The uploaded document contains no readable text.")

        document = models.Document(owner_id=user_id, filename=upload.filename, content=content)
        self.db.add(document)
        self.db.commit()
        self.db.refresh(document)

        try:
            self._create_chunks(document.id, document.filename, content)
            return document
        except Exception:
            # If chunking or embedding fails, remove the created document to keep DB consistent
            try:
                self.db.delete(document)
                self.db.commit()
            except Exception:
                # If cleanup fails, ignore here and re-raise original error
                pass
            raise

    def _extract_content(self, path: str, filename: str) -> str:
        if filename.lower().endswith(".pdf"):
            return self._read_pdf(path)
        elif filename.lower().endswith(".docx"):
            return self._read_docx(path)
        else:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()

    def _read_pdf(self, path: str) -> str:
        reader = PdfReader(path)
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(pages)

    def _read_docx(self, path: str) -> str:
        doc = DocxDocument(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        return "\n\n".join(paragraphs)

    def _create_chunks(self, document_id: int, filename: str, content: str) -> list[models.DocumentChunk]:
        chunks = split_text(content)
        if not chunks:
            return []

        chunk_models: list[models.DocumentChunk] = []
        batch_size = 10
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i : i + batch_size]
            # Prepend filename metadata to each chunk text for embedding and RAG context
            batch_with_metadata = [f"[Document: {filename}]\n{text}" for text in batch]
            attempts = 0
            while True:
                try:
                    embeddings = self.openai_service.embed_documents(batch_with_metadata)
                    if len(embeddings) != len(batch_with_metadata):
                        raise RuntimeError(f"Embedding API returned {len(embeddings)} embeddings for {len(batch_with_metadata)} inputs")

                    for idx, (chunk, embedding) in enumerate(zip(batch_with_metadata, embeddings)):
                        chunk_model = models.DocumentChunk(
                            document_id=document_id,
                            text=chunk,
                            embedding=embedding,
                            chunk_index=i + idx,
                        )
                        self.db.add(chunk_model)
                        chunk_models.append(chunk_model)

                    # Commit per batch to avoid huge transactions and make partial progress
                    self.db.commit()
                    break
                except Exception as exc:
                    attempts += 1
                    logger.exception("Failed to embed/insert chunk batch (attempt %s): %s", attempts, exc)
                    # rollback any pending DB changes for this session before retrying
                    try:
                        self.db.rollback()
                    except Exception:
                        pass
                    if attempts >= 3:
                        raise
                    # exponential backoff
                    time.sleep(2 ** attempts)

        return chunk_models
